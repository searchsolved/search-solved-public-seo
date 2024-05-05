"""Find Wikipedia citation opportunties. Lee Foot 05-05-2024
English only, for now!

More like this at: https://leefoot.co.uk
Contact: hello@leefoot.co.uk
"""

import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import requests
from bs4 import BeautifulSoup
import re
from io import BytesIO


def setup_streamlit():
    """
    Sets up the Streamlit page configuration with advanced styling for a centered layout.
    """
    st.set_page_config(page_title='Wikipedia Citation Finder V0.2 by Leefoot.co.uk', page_icon="🔗")
    st.markdown('# Wikipedia Citation Finder', unsafe_allow_html=True)
    st.caption("Find Wikipedia Pages Requiring Citations")
    st.markdown("""
        App by <a href="https://leefoot.co.uk" target="_blank">Lee Foot</a> | 
        Follow me on <a href="https://twitter.com/LeeFootSEO" target="_blank">Twitter</a> | 
        Need a bespoke app? <a href="mailto:hello@leefoot.co.uk">Get in Contact!</a>
    """, unsafe_allow_html=True)


def help_section():
    with st.expander("Help & Usage Guidelines"):
        st.write("""
        **Purpose of This App**
        The Wikipedia Citation Finder helps users identify Wikipedia articles that require citations. This tool is intended as a starting point for creating credible and authoritative content that can support Wikipedia's information needs.

        **Ethical Use Policy**
        It is crucial to use this tool responsibly. The intent is not to spam Wikipedia but to contribute valuable information that enhances the reliability and veracity of its content. Users are encouraged to:
        - Only create content relevant to their areas of expertise.
        - Ensure all contributions are backed by verifiable and reputable sources.
        - Aim to be a constructive part of the Wikipedia community by maintaining the integrity and quality of the information.

        **Creating Citeable Content**
        When using this app as inspiration to create content:
        - Choose topics where you have deep knowledge and where reputable sources can substantiate your contributions.
        - Avoid promotional language and focus on being educational.
        - Verify that your contributions provide real value to the respective topics and help address specific citation needs.
        """)


# Document Creation
def initialize_document():
    doc = Document()
    doc.add_heading('Citation Needed Report', 0)
    return doc

def add_summary_table(doc, results):
    doc.add_heading('Summary Table', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Wikipedia Page'
    hdr_cells[1].text = 'Number of Citations Needed'
    format_table_header(hdr_cells)
    return table

def format_table_header(hdr_cells):
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].font.bold = True

def fill_table(table, results):
    for url, citations in sorted(results.items(), key=lambda item: len(item[1]), reverse=True):
        row_cells = table.add_row().cells
        row_cells[0].text = url
        row_cells[1].text = str(len(citations))
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = paragraph.add_run(text)
    new_run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    new_run.font.underline = True
    hyperlink.append(new_run._r)
    paragraph._p.append(hyperlink)

def add_citations_to_document(doc, results):
    for url, citations in sorted(results.items(), key=lambda item: len(item[1]), reverse=True):
        doc.add_heading(url, level=2)
        for citation, section_url in citations:
            para = doc.add_paragraph(citation, style='ListBullet')  # Temporarily remove hyperlink to test

# Data Fetching and Processing
@st.cache_data
def get_wikipedia_urls(keyword):
    api_url = "https://en.wikipedia.org/w/api.php"
    params = {
        "action": "query",
        "list": "search",
        "srsearch": keyword,
        "format": "json"
    }
    response = requests.get(api_url, params=params)
    data = response.json()
    urls = [f"https://en.wikipedia.org/wiki/{item['title'].replace(' ', '_')}" for item in data["query"]["search"]]
    return urls

def extract_sentence(text, citation_tag):
    sentences = re.split(r'(?<=[.!?]) +', text)
    sentence = next((sentence for sentence in sentences if citation_tag in sentence), "Citation context not found.")
    section_url = text.split("#")[0].strip()
    return (sentence, section_url)

@st.cache_data
def search_citations_needed(urls):
    results = {}
    for url in urls:
        citations = find_citations(url)
        if citations:
            results[url] = citations
    return results

def find_citations(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    citations = soup.find_all(class_="noprint Inline-Template Template-Fact")
    return [extract_sentence(citation.find_parent('p').text if citation.find_parent('p') else '', citation.text) for citation in citations if citation.find_parent('p')]

# Streamlit UI and Main Function
def display_citation_report(citations_needed):
    if citations_needed:
        doc = initialize_document()
        table = add_summary_table(doc, citations_needed)
        fill_table(table, citations_needed)
        add_citations_to_document(doc, citations_needed)
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        st.download_button("Download Citation Report", data=doc_io, file_name="Citations_Needed_Report.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.error("No citations needed found for any of the pages.")

def main():
    setup_streamlit()
    keyword = st.text_input("Enter a keyword to search on Wikipedia:", "")
    if keyword:  # Check if the keyword is not empty
        urls = get_wikipedia_urls(keyword)
        if urls:  # Check if URLs are found
            citations_needed = search_citations_needed(urls)
            display_citation_report(citations_needed)
            for url, citations in citations_needed.items():
                with st.expander(f"{url} - {len(citations)} citations needed"):
                    st.markdown(f"[Visit page]({url})")
                    for citation, section_url in citations:
                        st.markdown(f"- {citation} [Link to section]({section_url})")
        else:
            st.error("No URLs found. Try a different keyword.")
    else:
        st.write("Enter a keyword above to search Wikipedia for pages needing citations.")

if __name__ == "__main__":
    main()
